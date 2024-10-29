import streamlit as st
import pandas as pd
import folium
from folium.plugins import MarkerCluster
import h3
from openai import OpenAI
from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO

# 제목과 부제 추가
st.title("SKT MEMS센서 메타데이터 관리")
st.subheader("자연어 명령어를 통해 손쉽게 MEMS를 검색하세요")

# OpenAI API 키 설정
client = OpenAI(api_key="sk-anudDrGdsMA3jiRirsA1erxqNcLA8r7ZtoVjazR6ltT3BlbkFJqEeof0VSlBHTNdzlQ3WZGLPWV520BTrBrPXQOB5K4A")

# 지도에 사용할 색상 함수
def get_marker_color(status):
    if status == 'normal':
        return 'green'
    elif status == 'disc.':
        return 'red'
    else:
        return 'blue'

# H3 경계를 지도에 그리는 함수
def draw_h3_boundaries(map_object, boundary_coords, total_count, normal_count, disc_count):
    folium.PolyLine(boundary_coords, color="darkred", weight=4, opacity=0.8).add_to(map_object)
    lat_center = sum([lat for lat, lon in boundary_coords]) / len(boundary_coords)
    lon_center = sum([lon for lat, lon in boundary_coords]) / len(boundary_coords)
    main_label = f'{total_count}'
    sub_label = f'({normal_count} normal / {disc_count} disc.)'
    folium.Marker(
        location=[lat_center, lon_center],
        popup=f"센서 수: {main_label} {sub_label}",
        icon=folium.DivIcon(html=f"""
            <div style="text-align: center;">
                <span style="font-size: 14pt; font-weight: bold; color: darkblue">{main_label}</span>
                <span style="font-size: 10pt; color: darkblue">{sub_label}</span>
            </div>
        """)
    ).add_to(map_object)

# 'df'가 세션 상태에 없으면 None으로 초기화
if 'df' not in st.session_state:
    st.session_state.df = None

# 파일 업로드
uploaded_file = st.file_uploader("최신 현황 파일을 업로드하세요", type=["xlsx", "xls", "csv"])
if uploaded_file is not None:
    try:
        if uploaded_file.name.endswith('.csv'):
            st.session_state.df = pd.read_csv(uploaded_file)
        else:
            st.session_state.df = pd.read_excel(uploaded_file, engine='openpyxl')

        # '단말번호'를 항상 11자리 문자열로 변환
        if '단말번호' in st.session_state.df.columns:
            st.session_state.df['단말번호'] = st.session_state.df['단말번호'].astype(str).str.zfill(11)

        st.subheader('업로드된 데이터')
        st.write(st.session_state.df)
    except Exception as e:
        st.error(f'파일 처리 중 오류 발생: {e}')

# '최신파일로 사용하기' 버튼
if st.button('최신파일로 사용하기'):
    try:
        st.write("최신 파일을 사용 중입니다.")
        st.session_state.df = pd.read_csv('Sensor_data_1024.csv')

        # '단말번호'를 항상 11자리 문자열로 변환
        if '단말번호' in st.session_state.df.columns:
            st.session_state.df['단말번호'] = st.session_state.df['단말번호'].astype(str).str.zfill(11)

        st.subheader('최신 파일 데이터')
        st.write(st.session_state.df.head())
    except Exception as e:
        st.error(f'최신 파일 불러오기 중 오류 발생: {e}')

# 현황 지도 보기 버튼
if st.session_state.df is not None:
    if st.button('현황 지도 보기'):
        if '위도' in st.session_state.df.columns and '경도' in st.session_state.df.columns and '연결상태' in st.session_state.df.columns:
            st.info("지도를 생성하고 있습니다...")

            try:
                # 지도 생성
                m = folium.Map(location=[st.session_state.df['위도'].mean(), st.session_state.df['경도'].mean()], zoom_start=12)

                # 마커 클러스터링 설정
                marker_cluster = MarkerCluster().add_to(m)

                # H3 인덱스별로 센서들을 클러스터링하기 위한 딕셔너리
                h3_dict = {}

                # 센서 데이터에 따라 H3 인덱스 계산 및 센서 개수 클러스터링
                for idx, row in st.session_state.df.iterrows():
                    lat, lon = row['위도'], row['경도']
                    h3_index = h3.latlng_to_cell(lat, lon, 5)
                    if h3_index not in h3_dict:
                        h3_dict[h3_index] = {
                            'total_count': 0,
                            'normal_count': 0,
                            'disc_count': 0,
                            'coords': h3.cell_to_boundary(h3_index)
                        }
                    h3_dict[h3_index]['total_count'] += 1
                    if row['연결상태'] == 'normal':
                        h3_dict[h3_index]['normal_count'] += 1
                    elif row['연결상태'] == 'disc.':
                        h3_dict[h3_index]['disc_count'] += 1

                    folium.Marker(
                        location=[lat, lon],
                        popup=f"연결상태: {row['연결상태']}",
                        icon=folium.Icon(color=get_marker_color(row['연결상태']))
                    ).add_to(marker_cluster)

                # H3 경계 및 센서 개수를 지도에 표시
                for h3_index, data in h3_dict.items():
                    boundary_coords = [(lat, lon) for lat, lon in data['coords']]
                    draw_h3_boundaries(
                        m,
                        boundary_coords,
                        data['total_count'],
                        data['normal_count'],
                        data['disc_count']
                    )

                # 지도를 HTML로 변환
                map_html = m._repr_html_()
                st.components.v1.html(map_html, width=700, height=500)
            except Exception as e:
                st.error(f"지도 생성 중 오류 발생: {e}")
        else:
            st.error("'위도', '경도', '연결상태' 컬럼이 데이터에 없습니다.")
else:
    st.info("데이터를 먼저 업로드하거나 최신 파일을 사용하세요.")

# 데이터 필터링을 위한 사용자 입력받기
st.subheader("데이터 필터링")

def generate_filter_condition(user_query, columns):
   try:
       response = client.chat.completions.create(
           model="gpt-3.5-turbo",
           messages=[
               {"role": "system", "content": "데이터프레임을 분석하여 필터링 조건을 생성하는 전문가입니다. 사용자의 명령어를 기반으로 조건을 만족하는 코드 구문을 생성하세요. 예를 들어, '경북 영덕군에 설치된 센서'라는 요청이 있으면 'df[df[\"설치지역\"] == \"경북 영덕군\"]' 형태로 반환하세요."},
               {"role": "user", "content": f"데이터프레임의 다음 컬럼으로 필터링할 조건을 만들어 주세요:\n\n{columns}\n\n사용자 요청: '{user_query}'\n\n주석이나 설명 없이 순수한 코드만 반환해 주세요."}
           ]
       )
       filter_code = response.choices[0].message.content.strip()
       filter_code_lines = filter_code.splitlines()
       executable_code = filter_code_lines[0].strip() if filter_code_lines else ""
       return executable_code
   except Exception as e:
       st.error(f"필터링 조건 생성 중 오류 발생: {e}")
       return None

# 데이터가 있는 경우에만 필터링 기능 실행
if st.session_state.df is not None:
    user_query = st.text_input("데이터 필터링을 위한 자연어 명령을 입력하세요")
    if user_query:
        filter_code = generate_filter_condition(user_query, st.session_state.df.columns.tolist())
        if filter_code:
            try:
                df = st.session_state.df
                filtered_df = eval(filter_code)
                st.subheader("필터링된 데이터")
                st.write(filtered_df)
                
               # 필터링된 데이터를 Word로 저장
                if st.button("필터링된 데이터를 MS Word로 저장"):
                    doc = Document()
                    doc.add_heading("필터링된 MEMS 센서 데이터", 0)
                    
                    # 예시 파일과 유사한 형식으로 데이터 추가
                    for idx, row in filtered_df.iterrows():
                    # 6열 x 10행 테이블 생성
                        table = doc.add_table(rows=10, cols=6)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 테이블을 페이지 가운데 정렬

                    
                    # 중앙 정렬 함수
                        def center_align(cell):
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 1행: 단말번호
                        table.cell(0, 0).text = "단말번호"
                        cell = table.cell(0, 1).merge(table.cell(0, 5))
                        cell.text = str(row.get("단말번호", ""))
                        center_align(table.cell(0, 0))
                        center_align(cell)
                        
                        # 2행: 관측소 코드, 시설구분, 시설구분세부
                        table.cell(1, 0).text = "관측소 코드"
                        table.cell(1, 1).text = str(row.get("관측소 코드", ""))
                        table.cell(1, 2).text = "시설구분"
                        table.cell(1, 3).text = str(row.get("시설구분", ""))
                        table.cell(1, 4).text = "시설구분세부"
                        table.cell(1, 5).text = str(row.get("시설구분세부", ""))
                        for col in range(6):
                            center_align(table.cell(1, col))
                        
                        # 3행: 제조사, 설치시점, 연결상태
                        table.cell(2, 0).text = "제조사"
                        table.cell(2, 1).text = str(row.get("제조사", ""))
                        table.cell(2, 2).text = "설치시점"
                        table.cell(2, 3).text = str(row.get("설치시점", ""))
                        table.cell(2, 4).text = "연결상태"
                        table.cell(2, 5).text = str(row.get("연결상태", ""))
                        for col in range(6):
                            center_align(table.cell(2, col))
                        
                        # 4행: 주소
                        table.cell(3, 0).text = "주소"
                        cell = table.cell(3, 1).merge(table.cell(3, 5))
                        cell.text = str(row.get("주소", ""))
                        center_align(table.cell(3, 0))
                        center_align(cell)
                        
                        # 5행: 위도, 경도, 고도
                        table.cell(4, 0).text = "위도"
                        table.cell(4, 1).text = str(row.get("위도", ""))
                        table.cell(4, 2).text = "경도"
                        table.cell(4, 3).text = str(row.get("경도", ""))
                        table.cell(4, 4).text = "고도"
                        table.cell(4, 5).text = str(row.get("고도", ""))
                        for col in range(6):
                            center_align(table.cell(4, col))
                        
                        # 6행: 설치층, 전체층, 축보정
                        table.cell(5, 0).text = "설치층"
                        table.cell(5, 1).text = str(row.get("설치층", ""))
                        table.cell(5, 2).text = "전체층"
                        table.cell(5, 3).text = str(row.get("전체층", ""))
                        table.cell(5, 4).text = "축보정"
                        table.cell(5, 5).text = str(row.get("축보정", ""))
                        for col in range(6):
                            center_align(table.cell(5, col))
                        
                        # 7행: H3 Cell, 센서 품질, 통신 품질
                        table.cell(6, 0).text = "H3 Cell"
                        table.cell(6, 1).text = str(row.get("H3 Cell", ""))
                        table.cell(6, 2).text = "센서 품질"
                        table.cell(6, 3).text = str(row.get("센서 품질", ""))
                        table.cell(6, 4).text = "통신 품질"
                        table.cell(6, 5).text = str(row.get("통신 품질", ""))
                        for col in range(6):
                            center_align(table.cell(6, col))
                        
                        # 8행: H3 혼잡여부, 센서 교체 필요 여부, 통신 품질 안정 여부
                        table.cell(7, 0).text = "H3 혼잡여부"
                        table.cell(7, 1).text = str(row.get("H3_Category", ""))
                        table.cell(7, 2).text = "센서 교체 필요 여부"
                        table.cell(7, 3).text = str(row.get("Sensor_Replacement_Status", ""))
                        table.cell(7, 4).text = "통신 품질 안정 여부"
                        table.cell(7, 5).text = str(row.get("Communication_Quality_Status", ""))
                        for col in range(6):
                            center_align(table.cell(7, col))
                        
                        # 9행: 현장 설치 사진
                        table.cell(8, 0).text = "현장 설치 사진"
                        cell = table.cell(8, 1).merge(table.cell(8, 5))
                        cell.text = "사진 링크: " + str(row.get("현장 설치 사진", ""))
                        center_align(table.cell(8, 0))
                        center_align(cell)
                        
                        # 10행: #Image_link1, #Image_link2
                        cell = table.cell(9, 0).merge(table.cell(9, 2))
                        cell.text = str(row.get("현장 설치 사진", "#Image_link1"))
                        center_align(cell)
                        
                        cell = table.cell(9, 3).merge(table.cell(9, 5))
                        cell.text = str(row.get("현장 설치 사진", "#Image_link2"))
                        center_align(cell)
                        
                        # 페이지 나누기
                        if idx < len(filtered_df) -1 :
                            doc.add_page_break()
                    
                    # 임시 버퍼에 저장
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    st.success("데이터가 MS Word 파일로 저장되었습니다.")
                    st.download_button("다운로드: 필터링된 데이터 Word 파일", data=buffer, file_name="filtered_sensor_data.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    
            except SyntaxError as se:
                st.error(f"구문 오류 발생: {se}")
            except Exception as e_eval:
                st.error(f"필터링 코드 실행 중 오류 발생: {e_eval}")
else:
    st.warning("데이터를 먼저 업로드해주세요.")
